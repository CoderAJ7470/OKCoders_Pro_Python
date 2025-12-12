# Build a data processing program

# Read data from a csv file, then store that in an SQLite DB
# Query the DB for analysis (top grossing movies, top audience scores, etc)
# Generate a Word document report with tables
# Export the results to JSON format

import csv, sqlite3, json
from docx import Document

def read_csv_file_into_db():
    movie_file_path = 'C:/OKCoders_Pro_Python/movies.csv'

    # Read data from the csv file
    movie_file = open(movie_file_path)
    movie_file_reader = csv.reader(movie_file)
    movie_data_list = list(movie_file_reader)

    store_movie_data_in_sqlitedb(movie_data_list)

def store_movie_data_in_sqlitedb(data_list):
    # isolation_level specifies how changes to the db are committed. In this case, specifying None tells sqlite to auto-commit each transaction, so we do not have to manually call commit() each time we want to save data in the db or update something
    movie_database = sqlite3.connect('movie_data.db', isolation_level = None)
    
    movie_database.execute(f'CREATE TABLE IF NOT EXISTS movie_data (Film TEXT NOT NULL, Genre TEXT, Lead_Studio TEXT, Audience_Score_Pct TEXT, Profitability TEXT, Rotten_Tomatoes_Pct TEXT, Worldwide_Gross TEXT, Release_Year TEXT)')

    # Get the data for each movie, from the data_list that is passed in, starting art the list at index 1. So here, data gets a list of nested lists, beginning at index 1 in the data_list coming in
    data = data_list[1:]

    # Insert the data into the db. The function executemany() allows us to specify a parameterized query for each nested list in data above. When the function runs with the given arguments, it will automatically loop through all nested lists and input the data from them as specified by the question marks. Remember to close the connection once done.
    movie_database.executemany('INSERT INTO movie_data VALUES (?, ?, ?, ?, ?, ?, ?, ?)', data)
    
    movie_database.execute('DELETE FROM movie_data WHERE rowid NOT IN (SELECT MIN(rowid) FROM movie_data GROUP BY Film);')

    movie_database.close()
    
# Getter functions (return data only)
def get_top_ten_grossing_movies_worldwide():
    movie_db = sqlite3.connect('movie_data.db', isolation_level = None)
    
    query = '''
        SELECT Film, Genre, Worldwide_Gross
        FROM movie_data
        ORDER BY CAST(REPLACE(Worldwide_Gross, '$', '') AS REAL) DESC
        LIMIT 10;
    '''
    
    rows = movie_db.execute(query).fetchall()
    movie_db.close()
    return rows

def get_top_ten_profitable_movies():
    movie_db = sqlite3.connect('movie_data.db', isolation_level = None)
    
    query = '''
        SELECT Film, Genre, Profitability
        FROM movie_data
        ORDER BY CAST(Profitability AS REAL) DESC
        LIMIT 10;
    '''
    
    rows = movie_db.execute(query).fetchall()
    movie_db.close()
    return rows

def get_movies_by_top_ten_audience_percentages():
    movie_db = sqlite3.connect('movie_data.db', isolation_level = None)
    
    query = '''
        SELECT Film, Genre, Audience_Score_Pct
        FROM movie_data
        ORDER BY CAST(Audience_Score_Pct AS INTEGER) DESC
        LIMIT 10;
    '''
    
    rows = movie_db.execute(query).fetchall()
    movie_db.close()
    return rows

# Functions that output to the terminal
def show_top_ten_grossing_movies_worldwide():
    rows = get_top_ten_grossing_movies_worldwide()
    
    print('\nShowing top ten worldwide-grossing movies from 2007-2011, in descending order:\n')
    for index, (film, genre, gross_profit) in enumerate(rows, start = 1):
        print(f'{index}. {film} ({genre}) — {gross_profit} million')

def show_top_ten_profitable_movies():
    rows = get_top_ten_profitable_movies()
    
    print('\nShowing top ten profitable movies from 2007-2011, in descending order. The higher the profitability ratio, the more profit the movie made:\n')
    for index, (film, genre, profit) in enumerate(rows, start = 1):
        print(f'{index}. {film} ({genre}) — Profitability: {float(profit):.2f}')

def movies_by_top_ten_audience_percentages():
    rows = get_movies_by_top_ten_audience_percentages()
    
    print('Showing movies with the top 10 audience score percentages:\n')
    for index, (film ,genre, audience_percentage) in enumerate(rows, start = 1):
        print(f'{index}. {film} ({genre}) - {audience_percentage}%')

def export_report_results_to_json():
    # Collect results from the getter functions
    grossing_movies = get_top_ten_grossing_movies_worldwide()
    profitable_movies = get_top_ten_profitable_movies()
    audience_movies = get_movies_by_top_ten_audience_percentages()

    # Build a dictionary to hold all results
    report_data = {
        "Top 10 Worldwide Grossing Movies": [
            {"Film": film, "Genre": genre, "Worldwide_Gross": gross}
            for film, genre, gross in grossing_movies
        ],
        "Top 10 Profitable Movies": [
            {"Film": film, "Genre": genre, "Profitability": profit}
            for film, genre, profit in profitable_movies
        ],
        "Top 10 Audience Score Percentages": [
            {"Film": film, "Genre": genre, "Audience_Score_Pct": score}
            for film, genre, score in audience_movies
        ]
    }

    # Export to JSON file
    with open("top_ten_data.json", "w", encoding = "utf-8") as json_file:
        json.dump(report_data, json_file, indent = 4)

    print("JSON report generated: top_ten_data.json")
    
def print_movie_db():
    movie_db = sqlite3.connect('movie_data.db', isolation_level = None)
    rows = movie_db.execute('SELECT * FROM movie_data').fetchall()
   
    for index, row in enumerate(rows, start = 1):
        print(f'{index}.', row) # prints each tuple (row) in the db
        
    movie_db.close()
    
# Generates a Word file with all of the "top 10" data in tabular format
def generate_word_doc_report():
    doc = Document()
    doc.add_heading('Movie Data Reports (Top 10)', level = 0)

    # Section 1: Top 10 Worldwide Grossing Movies
    doc.add_heading('Top 10 Worldwide Grossing Movies', level = 1)
    grossing_movies = get_top_ten_grossing_movies_worldwide()
    
    table = doc.add_table(rows = 1, cols = 3)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Film'
    header_cells[1].text = 'Genre'
    header_cells[2].text = 'Worldwide Gross'
    
    for film, genre, gross in grossing_movies:
        row_cells = table.add_row().cells
        row_cells[0].text = film
        row_cells[1].text = genre
        row_cells[2].text = gross

    # Section 2: Top 10 Profitable Movies
    doc.add_heading('Top 10 Profitable Movies', level = 1)
    profitable_movies = get_top_ten_profitable_movies()
    
    table = doc.add_table(rows = 1, cols = 3)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Film'
    header_cells[1].text = 'Genre'
    header_cells[2].text = 'Profitability'
    
    for film, genre, profit in profitable_movies:
        row_cells = table.add_row().cells
        row_cells[0].text = film
        row_cells[1].text = genre
        row_cells[2].text = str(profit)

    # Section 3: Top 10 Audience Score Percentages
    doc.add_heading('Top 10 Audience Score Percentages', level = 1)
    audience_movies = get_movies_by_top_ten_audience_percentages()
    
    table = doc.add_table(rows = 1, cols = 3)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Film'
    header_cells[1].text = 'Genre'
    header_cells[2].text = 'Audience Score %'
    
    for film, genre, score in audience_movies:
        row_cells = table.add_row().cells
        row_cells[0].text = film
        row_cells[1].text = genre
        row_cells[2].text = str(score)

    doc.save('Top_10_Reports.docx')
    print("Word report generated: Top_10_Reports.docx")
    
def show_user_options():
    # Storing menu option in this dictionary; allow for additional choices to be added later if required
    menu_options = {
        '1': ('Show top 10 worldwide-grossing movies', show_top_ten_grossing_movies_worldwide),
        '2': ('Show top 10 profitable movies', show_top_ten_profitable_movies),
        '3': ('Show movies by top 10 audience percentages', movies_by_top_ten_audience_percentages),
        '4': ('Print all movie data', print_movie_db),
        '5': ('Generate Word report for all Top 10 data', generate_word_doc_report),
        '6': ('Export results to JSON', export_report_results_to_json)
    }
    
    print('\nWelcome! This program allows you to load a movie database and perform select actions with that data. Enter "q" anytime to quit.\n')

    while True:
        # Show menu options
        for key, (label, _) in menu_options.items():
            print(f'{key}. {label}')

        choice = input("\nEnter your choice here: ").strip()

        if choice.lower() == "q":
            break

        # Validate menu choice
        if choice not in menu_options:
            print("Invalid option — please choose a valid number.")
            continue
        
        # menu_options[choice] returns a tuple, from which you can select the label and the function, based on which option (choice) the user selected
        selected_label, selected_function = menu_options[choice]

        print(f'\nYou selected: {selected_label}\n')
        selected_function() # Run the function selected by the user/chosen from the menu
        
        print('\nKeep choosing, or q to quit:\n')

if __name__ == '__main__':
    show_user_options()
